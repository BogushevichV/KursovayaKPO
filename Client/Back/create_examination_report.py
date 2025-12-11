from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from PySide6.QtWidgets import QMessageBox
from Client.Back.report_manager import ReportManager
from Client.Source.config import SERVER_URL


class CreateExaminationReport:
    @staticmethod
    def create_report(server_url=None, form_data=None, filename="exam_report.docx"):
        """Создание ведомости через сервер API"""
        try:
            if server_url is None:
                server_url = SERVER_URL

            if form_data is None:
                raise ValueError("form_data обязателен")

            # Инициализируем ReportManager для работы через сервер
            report_manager = ReportManager(server_url=server_url)

            # Получаем данные о студентах и оценках через сервер
            group_number = form_data.get('group', '')
            subject_name = form_data.get('subject', '')
            course = form_data.get('course', '')
            semester = form_data.get('semester', '')

            # Получаем студентов и оценки через API сервера
            students_data = report_manager.find_subject_grades(
                subject_name=subject_name,
                group_number=group_number,
                course=course,
                semester=semester
            )

            # Преобразуем данные в нужный формат
            students = []
            if students_data:
                for item in students_data:
                    if isinstance(item, (list, tuple)) and len(item) >= 3:
                        students.append((item[0], item[1], item[2]))
                    elif isinstance(item, dict):
                        students.append((
                            item.get('name', ''),
                            item.get('gradebook', ''),
                            item.get('grade', '')
                        ))

            # Вычисляем статистику по оценкам из полученных данных
            grade_stats = {}
            for _, _, grade in students:
                if grade:
                    # Обрабатываем оценки в баллах (0-10)
                    try:
                        grade_num = int(float(grade))
                        if 0 <= grade_num <= 10:
                            grade_stats[str(grade_num)] = grade_stats.get(str(grade_num), 0) + 1
                    except (ValueError, TypeError):
                        # Если не число, пропускаем
                        pass

            # Подсчитываем количество студентов с отметками (явившихся)
            present_count = len([s for s in students if s[2] and s[2].strip()])

            # Количество студентов без отметок (не явившихся)
            absent_count = len([s for s in students if not s[2] or not s[2].strip()])

            # Создаем документ Word
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # Устанавливаем межстрочный интервал 0
            paragraph_format = style.paragraph_format
            paragraph_format.space_after = Pt(0)
            paragraph_format.space_before = Pt(0)

            # Заголовок
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run("БЕЛОРУССКИЙ НАЦИОНАЛЬНЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ")
            title_run.bold = True
            title_run.font.size = Pt(14)

            doc.add_paragraph()  # Пустая строка

            # Подзаголовок
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Добавляем первую часть текста без подчеркивания
            subtitle_run1 = subtitle.add_run("ЗАЧЁТНО-ЭКЗАМЕНАЦИОННАЯ ВЕДОМОСТЬ № ")
            subtitle_run1.bold = True
            subtitle_run1.font.size = Pt(12)

            # Добавляем номер ведомости с подчеркиванием
            subtitle_run2 = subtitle.add_run(form_data['statement_number'])
            subtitle_run2.bold = True
            subtitle_run2.font.size = Pt(12)
            subtitle_run2.underline = True

            # Строка промежуточной аттестации
            intermediate = doc.add_paragraph()
            intermediate.alignment = WD_ALIGN_PARAGRAPH.CENTER
            intermediate_run = intermediate.add_run("промежуточной аттестации учебной группы")
            intermediate_run.bold = True
            intermediate.paragraph_format.space_after = Pt(0)

            doc.add_paragraph()

            # Основная информация (подчеркиваем только значения из формы)
            def add_paragraph_with_underline(label, value):
                p = doc.add_paragraph()
                p.add_run(label)
                run = p.add_run(str(value))
                run.underline = True
                p.paragraph_format.space_after = Pt(0)

            add_paragraph_with_underline("Вид высшего образования: ", form_data['education_type'])
            add_paragraph_with_underline("Форма получения высшего образования: ", form_data['study_form'])
            add_paragraph_with_underline("Форма промежуточной аттестации: ", form_data['exam_type'])
            year_semester_p = doc.add_paragraph()
            year_semester_p.add_run("Учебный год ")
            year_run = year_semester_p.add_run(str(form_data['year']))
            year_run.underline = True
            year_semester_p.add_run(" Семестр ")
            semester_run = year_semester_p.add_run(str(form_data['semester']))
            semester_run.underline = True
            year_semester_p.paragraph_format.space_after = Pt(0)

            add_paragraph_with_underline("Факультет ", form_data['faculty'])
            p = doc.add_paragraph()
            p.add_run("Курс ")
            run = p.add_run(str(form_data['course']))
            run.underline = True
            p.add_run(" группа ")
            run = p.add_run(form_data['group'])
            run.underline = True
            p.paragraph_format.space_after = Pt(0)

            add_paragraph_with_underline("Название учебной дисциплины, модуля, практики ", form_data['subject'])
            add_paragraph_with_underline("Всего часов по учебной дисциплине, модулю, практике в семестре ",
                                         form_data['hours'])
            add_paragraph_with_underline("Зачетных единиц по учебной дисциплине, модулю, практике в семестре ",
                                         f"{form_data['credits']} з.е.")
            add_paragraph_with_underline("Фамилия, инициалы преподавателя(ей) ", form_data['teacher'])
            add_paragraph_with_underline("Дата проведения аттестации ", form_data['exam_date'])
            add_paragraph_with_underline("Формат проведения аттестации ", form_data['exam_format'])

            doc.add_paragraph()

            # Таблица студентов (6 колонок)
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'

            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            headers = ['№ п/п', 'Фамилия, инициалы', '№ зачетной книжки',
                       'Отметка о зачете', 'Отметка в баллах', 'Подпись преподавателя(ей)']

            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Увеличиваем ширину колонки для ФИО
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    if idx == 1:  # Колонка с ФИО
                        cell.width = Pt(180)
                    else:
                        cell.width = Pt(55)

            # Заполняем таблицу данными из БД
            for i, (name, gradebook, grade) in enumerate(students, 1):
                row_cells = table.add_row().cells

                # № п/п
                row_cells[0].text = str(i)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # ФИО
                row_cells[1].text = name
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # № зачетной книжки
                row_cells[2].text = gradebook
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Отметка о зачете (выравнивание по центру)
                if grade and grade.lower() in ['зачтено', 'не зачтено']:
                    row_cells[3].text = grade
                    row_cells[4].text = ""
                else:
                    row_cells[3].text = ""
                    if grade:
                        row_cells[4].text = str(grade)
                    else:
                        row_cells[4].text = ""

                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Отметка в баллах (выравнивание по центру)
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Подпись преподавателя (пустая)
                row_cells[5].text = ""
                row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Подпись декана после таблицы
            doc.add_paragraph()  # Пустая строка

            # Подпись декана
            dean_title = doc.add_paragraph()
            dean_title.add_run("Декан                                                                                "
                               "                    ")
            dean_run = dean_title.add_run(form_data['dean'])
            dean_run.underline = True
            dean_title.paragraph_format.space_after = Pt(0)

            # Добавляем подпись в скобках с новой строки
            director_line = doc.add_paragraph()
            director_run = director_line.add_run("(директор института, филиала)                             "
                                                 "            подпись, м.п.                                     "
                                                 "инициалы, фамилия")
            director_line.paragraph_format.space_after = Pt(0)
            director_run.font.size = Pt(9)

            doc.add_paragraph()  # Пустая строка в конце

            # Статистика
            present_p = doc.add_paragraph()
            present_p.add_run("Количество обучающихся, присутствующих на аттестации: ")
            present_run = present_p.add_run(str(present_count))
            present_run.underline = True
            present_p.paragraph_format.space_after = Pt(0)

            doc.add_paragraph("Количество обучающихся, получивших отметки:")
            doc.paragraphs[-1].paragraph_format.space_after = Pt(0)

            # Форматирование статистики как в примере
            grade_lines = [
                ['10 (десять)', '8 (восемь)', '5 (пять)', '3 (три)'],
                ['9 (девять)', '7 (семь)', '4 (четыре)', '2 (два)'],
                ['6 (шесть)', '1 (один)']
            ]

            for line in grade_lines:
                line_p = doc.add_paragraph()
                for i, grade in enumerate(line):
                    grade_num = grade.split()[0]
                    count = grade_stats.get(grade_num, 0)

                    # Для цифры 6 добавляем дополнительные пробелы для выравнивания
                    if grade_num == '6':
                        line_p.add_run("                                 ")

                    line_p.add_run(f"{grade}___")
                    count_run = line_p.add_run(str(count))
                    count_run.underline = True
                    line_p.add_run("___")
                line_p.paragraph_format.space_after = Pt(0)

            absent_p = doc.add_paragraph()
            absent_p.add_run("Количество обучающихся, не явившихся на аттестацию ")
            absent_p.add_run("(в том числе не допущенных к аттестации): ")
            absent_run = absent_p.add_run(str(absent_count))
            absent_run.underline = True
            absent_p.paragraph_format.space_after = Pt(0)

            # Сохраняем документ
            doc.save(filename)
            return filename

        except Exception as e:
            QMessageBox.critical(None, "Ошибка", f"Ошибка при создании ведомости: {str(e)}")
            return None

